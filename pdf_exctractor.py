from fastapi import FastAPI, UploadFile, File, Request,Form, Response
import tempfile, os, json, re
from docx import Document
import pdfplumber, requests, traceback
from bs4 import BeautifulSoup
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from typing import Any, Dict
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Optional

# Upstream LLM endpoint (unchanged)
LLM_API_URL = "https://file-extractordev.sidbi.in/fill-docx-service/extract"
# Optional: upstream timeout seconds
UPSTREAM_TIMEOUT = int(os.getenv("UPSTREAM_TIMEOUT", "240"))

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins =[
        "*"
    ],
    allow_credentials = True,
    allow_methods = ["*"],
    allow_headers = "*",
)

def to_crores(val):
    if val is None:
        return "Couldn't be Captured"
    if isinstance(val, str):
        v = val.strip()
        if v == "" or v == "NA":
            return "Couldn't be Captured"
        low = v.lower().replace(",", "")
        nums = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*lakhs?", low)
        if nums:
            num = float(nums.group(1))
            return f"{round(num/100,2)}"
        return v
    return str(val)

class HtmlToDocxPayload(BaseModel):
    html: str
    file_name: Optional[str] = "Output.docx"

class FillHtmlPayload(BaseModel):
    data:Dict[str,Any]
    template_html:str

def get_cell_text(cell):
    t = cell.text
    if t is None: return None
    t = t.strip()
    return t if t else None

def cell_merged(cell):
    return getattr(cell._tc, "gridSpan", 1) > 1

def is_fully_merged_row(row):
    spans = [getattr(c._tc, "gridSpan", 1) for c in row.cells]
    return max(spans) > 1 and len(set(spans)) == 1

def is_srno(x):
    if x is None: return False
    x2 = x.strip().lower().replace('.', '').replace(' ', '').replace('\n','')
    if x2.isdigit(): return True
    return x2 in ["srno", "sr","sno"]

def clean_keys(obj):
    if isinstance(obj,dict):
        new = {}
        for k,v in obj.items():
            k = k.replace("\n"," ").strip()
            k = "_".join(k.split())
            new[k] = clean_keys(v)
        return new
    if isinstance(obj,list):
        return[clean_keys(i) for i in obj]
    if isinstance(obj,str):
        obj = obj.replace("\n"," ").strip()
        obj = "_".join(obj.split())
        return obj
    else:
        obj = obj.replace(" ","_")
    return obj

def extract_docx_json(path):
    doc = Document(path)
    results = []

    def combine_headers(headers, idx=0, prefix=""):
        if idx >= len(headers): return [prefix]
        row = headers[idx]
        result = []
        for col in row:
            col_clean = col.strip() if col else ""
            new_prefix = f"{prefix}_{col_clean}" if prefix else col_clean
            result.extend(combine_headers(headers, idx+1, new_prefix))
        return result

    for el in doc.element.body:
        if el.tag.endswith('tbl'):
            tbl = None
            for t in doc.tables:
                if t._element == el:
                    tbl = t
                    break
            if tbl is None: continue
            rows_text = [[get_cell_text(c) for c in r.cells] for r in tbl.rows]
            if not rows_text: continue

            has_partial = any(
                any(v is None or v.strip() == "" for v in r) and
                any(v is not None and v.strip() != "" for v in r)
                for r in rows_text
            )

            cleaned_rows = []
            for r in rows_text:
                if all(v and v.strip() for v in r):
                    if has_partial:
                        cleaned_rows.append(r[1:])
                    else:
                        cleaned_rows.append(r)
                else:
                    cleaned_rows.append(r)
            rows_text = cleaned_rows

            first_cell = rows_text[0][0] if rows_text[0] else None
            ignore_first_col = is_srno(first_cell)

            processed_rows = []
            first_none_idx = None
            for idx, r in enumerate(rows_text):
                if any(v is None or v.strip() == "" for v in r):
                    first_none_idx = idx
                    break

            for idx, (r, orig_row) in enumerate(zip(rows_text, tbl.rows)):
                merged_full = is_fully_merged_row(orig_row)
                merged_any = any(cell_merged(c) for c in orig_row.cells)
                if merged_full or merged_any:
                    if idx != 0:
                        continue
                processed_rows.append(r)

            if first_none_idx is not None:
                filtered_rows = []
                for idx, r in enumerate(processed_rows):
                    if first_none_idx < idx:
                        if all(v and v.strip() for v in r):
                            continue
                    filtered_rows.append(r)
                processed_rows = filtered_rows
            
            adjusted_rows = []
            for r in processed_rows:
                fully_filled = all(v and v.strip() for v in r)
                if fully_filled and ignore_first_col and len(r) > 1:
                    r = r[1:]
                adjusted_rows.append(r)
            processed_rows = adjusted_rows

            column_hierarchy = []
            row_headers = []

            for idx, r in enumerate(processed_rows):
                fully_filled = all(v and v.strip() for v in r)
                if fully_filled and (first_none_idx is None or idx < first_none_idx):
                    column_hierarchy.append(r)
                else:
                    cleaned = [v for v in r[:2] if v and v.strip()]
                    if cleaned:
                        if len(cleaned) == 1:
                            row_headers.append(cleaned[0])
                        else:
                            non_digits = [x for x in cleaned if not x.isdigit()]
                            if non_digits:
                                row_headers.append(non_digits[0])
                            else:
                                row_headers.append(cleaned[0])
        
            row_headers = clean_keys(row_headers)
            column_hierarchy = clean_keys(column_hierarchy)
            
            main_row_col_idx = None
            if row_headers and column_hierarchy:
                for r in processed_rows:
                    for i, v in enumerate(r):
                        if v and v.strip() and clean_keys(v) in row_headers: #type: ignore
                            main_row_col_idx = i
                            break
                    if main_row_col_idx is not None:
                        break
                if main_row_col_idx is not None:
                    column_hierarchy = [r[main_row_col_idx:] for r in column_hierarchy]

            flat_col_headers = []
            for r in column_hierarchy:
                for c in r:
                    if c not in flat_col_headers:
                        flat_col_headers.append(c)

            placeholders = []
            if column_hierarchy and row_headers:
                col_paths = combine_headers(column_hierarchy)
                for cp in row_headers:
                    for rh in col_paths:
                        placeholders.append(cp + "_" + rh)
            elif column_hierarchy:
                col_paths = combine_headers(column_hierarchy)
                for cp in flat_col_headers:
                    placeholders.append(cp)
            elif row_headers:
                for rh in row_headers:
                    placeholders.append(rh)

            placeholders = list(dict.fromkeys(placeholders))

            placeholders = list(dict.fromkeys(placeholders))

            results.append({
                "placeholders": {th: "" for th in placeholders}
            })

        elif el.tag.endswith('p'):
            p_text = ''.join([node.text for node in el if node.text])
            matches = re.findall(r'\[(.*?)\]', p_text)
            for m in matches:
                results.append({m.strip().replace(" ","_"): ""})

    return results

def normalize_key(k):
    k = k.strip()
    k = k.lower()
    k = re.sub(r'\s+', '_', k)
    k = re.sub(r'[^a-z0-9_]', '_', k)
    k = re.sub(r'_+', '_', k)
    k = k.strip('_')
    return k

def fill_html(html_content: str, json_content: str) -> str:
    data = json.loads(json_content).get("extracted_json", [])

    placeholders_list = [d["placeholders"] for d in data if "placeholders" in d]
    non_placeholder_data = {}
    for d in data:
        if "placeholders" not in d:
            for k, v in d.items():
                nk = normalize_key(k)
                non_placeholder_data[nk] = v

    soup = BeautifulSoup(html_content, "html.parser")

    tables = soup.find_all("table")
    for t_idx, table in enumerate(tables):
        if t_idx >= len(placeholders_list):
            break
        placeholders = list(placeholders_list[t_idx].values())
        p_idx = 0
        for tr in table.find_all("tr"):
            cells = tr.find_all(["th","td"])
            if not cells or not cells[0].get_text(strip=True):
                continue
            if any(not cell.get_text(strip=True) and p_idx < len(placeholders) for cell in cells):
                for cell in cells:
                    if not cell.get_text(strip=True) and p_idx < len(placeholders):
                        cell.string = placeholders[p_idx]
                        p_idx += 1
            if p_idx >= len(placeholders):
                break

    for p in soup.find_all("p"):
        text = p.get_text()
        matches = re.findall(r'\[(.*?)\]', text)
        if matches:
            new_text = text
            for m in matches:
                key = normalize_key(m)
                if key in non_placeholder_data:
                    new_text = new_text.replace(f"[{m}]", str(non_placeholder_data[key]))
            p.string = new_text

    return str(soup)


def fill_html2(html_content: str, json_content: str) -> str:
    try:
        data = json.loads(json_content).get("extracted_json", [])

        table_data_list=[list(d["placeholders"].values()) for d in data if "placeholders" in d]
        text_data=[]
        for d in data:
            if "placeholders" not in d:
                for _,v in d.items(): 
                    text_data.append(v)
        soup=BeautifulSoup(html_content,"html.parser")
        for tag in soup.find_all(["p","li","span","strong","td","th"]):
            tag.string = tag.get_text()
        soup.smooth()
        for tag in soup.find_all(["p","li","span","strong","td","th"]):
            if tag.name=="p": tag["style"]="font-family:Rupee Foradian;font-size:11pt"
            else: tag["style"]="font-family:Rupee Foradian;font-size:10pt"
        tables=soup.find_all("table")
        for t_idx,table in enumerate(tables):
            if t_idx>=len(table_data_list): break
            values=table_data_list[t_idx]
            vi=0
            for tr in table.find_all("tr"):
                cells=tr.find_all(["td","th"])
                if not cells: 
                    continue
                texts=[c.get_text(strip=True) for c in cells]
                if len(texts) == 1: continue
                if len(texts)>=2 and texts[0]=="" and texts[1]=="": 
                    continue
                start_idx=0
                if len(texts)>=2 and texts[0]=="" and texts[1]!="": 
                    start_idx=1
                for i in range(start_idx,len(cells)):
                    if vi<len(values) and texts[i]=="": cells[i].string=to_crores(values[vi]); vi+=1
        
        text_idx = 0
        
        for i in range(len(text_data)):
            if text_data[i] == "" or text_data[i].lower() == "na":
                text_data[i] = "Couldn't be Captured"
        
        for node in soup.descendants:
            if not isinstance(node, NavigableString):
                continue
            
            parent = node.parent
            if parent and parent.name == "table":
                continue
            
            text = str(node)
            if "[" not in text or "]" not in text:
                continue
            
            parts = re.split(r'(\[[^\[\]]*\])', text)
            new_text = ""
            
            for part in parts:
                if part.startswith("[") and part.endswith("]"):
                    if text_idx< len(text_data):
                        new_text += text_data[text_idx]
                        text_idx += 1
                    else:
                        new_text += "Couldn't be Captured."
                else:
                    new_text += part
                    
            if new_text is None or new_text == "" or new_text.lower() == "na":
                new_text = "Couldn't be Captured."
        
        
            node.replace_with(NavigableString(new_text))
            
        return str(soup)


    except Exception as e:
        return f"<html><body><h3>Error processing HTML: {e}</h3></body></html>"


@app.post("/llm-extract")
async def proxy_llm_extract(request: Request):
    """
    Accept multipart/form-data (txt_file + input_schema etc.) from browser,
    forward to the upstream LLM API, and return the upstream response as-is
    (status code + content-type + body).
    """
    try:
        # Read the incoming form data (starve memory for very huge files in prod)
        form = await request.form()

        # Build files and data dicts for requests.post
        files = {}
        data = {}

        # form.multi_items() works well to iterate form-data pairs
        for key, value in form.multi_items():
            # If value is UploadFile-like, forward as file
            if hasattr(value, "filename") and value.filename:
                # value.file is a SpooledTemporaryFile; read bytes
                file_content = value.file.read()
                content_type = getattr(value, "content_type", "application/octet-stream")
                files[key] = (value.filename, file_content, content_type)
            else:
                # normal form field
                data[key] = value

        # Forward to upstream LLM
        resp = requests.post(
            LLM_API_URL,
            files=files if files else None,
            data=data if data else None,
            timeout = UPSTREAM_TIMEOUT
        )

        # Build response to client with the same content-type and status
        content_type = resp.headers.get("Content-Type", None)
        return Response(content=resp.content, status_code=resp.status_code, media_type=content_type)

    except requests.RequestException as e:
        # Upstream request failed (network/timeout/connection)
        traceback.print_exc()
        return JSONResponse(
            status_code=502,
            content={"detail": f"Upstream LLM request failed: {str(e)}"}
        )
    except Exception as e:
        # Generic server error
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"detail": f"Internal server error: {str(e)}"}
        )


@app.post("/fill-html", response_class=HTMLResponse)
async def fill_html_api(request: Request):
    form = await request.form()
    html_file = None
    json_text = None

    for key, value in form.items():
        if hasattr(value, "filename"):  # file
            html_file = value
        else:  # text
            json_text = value

    if not html_file or not json_text:
        return HTMLResponse(content="HTML file and JSON text are required", status_code=400)

    html_content = await html_file.read() # type: ignore
    filled_html = fill_html(html_content.decode("utf-8"), json_text) # type: ignore
    return HTMLResponse(content=filled_html)



@app.post("/fill-html2", response_class=HTMLResponse)
async def fill_html2_api(payload:FillHtmlPayload):
    try:
        html_content = payload.template_html
        json_text = json.dumps(payload.data)


        filled_html = fill_html2(html_content, json_text)
        return HTMLResponse(content=filled_html)
    except Exception as e:
        print("Error in /fill-html:", str(e))
        return HTMLResponse(content=f"Internal Server Error: {e}", status_code=500)


@app.post("/docx-to-json")
async def docx_to_json(request: Request):
    form  = await request.form()
    file = None
    for v in form.values():
        if hasattr(v,"filename") or hasattr(v,'file'):
            file = v
            break
    if file is None:
        return{"Error" :"No file provided"}
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(await file.read()) # type: ignore
        tmp_path = tmp.name
    out = extract_docx_json(tmp_path)
    os.remove(tmp_path)
    return out

# @app.post("/pdf-to-text")
# async def pdf_to_text(request: Request):
#     form  = await request.form()
#     file = None
#     for v in form.values():
#         if hasattr(v,"filename") or hasattr(v,'file'):
#             file = v
#             break
#     if file is None:
#         return{"Error" :"No file provided"}    
#     with tempfile.NamedTemporaryFile(delete=False) as tmp:
#         tmp.write(await file.read()) # type: ignore
#         tmp_path = tmp.name
#     text = ""
#     with pdfplumber.open(tmp_path) as pdf:
#         for p in pdf.pages:
#             t = p.extract_text()
#             if t: text += t + "\n"
#     os.remove(tmp_path)
#     return text

import time, pymupdf
@app.post("/pdf-to-text")
async def pdf_to_text(request: Request):
    """
    Upload a PDF via multipart/form-data and extract text using PyMuPDF.

    Uncomment ONE of the two extraction blocks below:
    1) LINE-BY-LINE (simple linear page text)
    2) BLOCK-ORDERED (sorted by y0 then x0 for better reading order)
    """

    # Parse multipart/form-data and find the uploaded file
    try:
        form = await request.form()
    except Exception as e:
        return JSONResponse({"Error": f"Invalid form-data: {e}"}, status_code=400)

    file = None
    for v in form.values():
        if hasattr(v, "filename") or hasattr(v, "file"):
            file = v
            break

    if file is None:
        return JSONResponse({"Error": "No file provided"}, status_code=400)

    # Read file bytes from the upload (avoid temp files)
    try:
        pdf_bytes = await file.read()
    except Exception as e:
        return JSONResponse({"Error": f"Failed to read uploaded file: {e}"}, status_code=400)

    if not pdf_bytes:
        return JSONResponse({"Error": "Uploaded file is empty"}, status_code=400)

    # Open PDF via PyMuPDF (from bytes)
    start = time.time()
    try:
        doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        return JSONResponse({"Error": f"Failed to open PDF: {e}"}, status_code=400)

    parts = []
    try:
        # ----------------------------------------------------------
        # OPTION 1: LINE-BY-LINE (simple linearized text per page)
        # ----------------------------------------------------------
        # for page in doc:
        #     t = page.get_text()  # linearized text
        #     if t:
        #         t = t.strip()
        #         if t:
        #             parts.append(t)

        # ----------------------------------------------------------
        # OPTION 2: BLOCK-ORDERED (better reading order for columns)
        # ----------------------------------------------------------
        for page in doc:
            blocks = page.get_text("blocks")
            # Each block: (x0, y0, x1, y1, text, block_no, ...)
            blocks.sort(key=lambda b: (b[1], b[0]))  # sort by top, then left
            for b in blocks:
                text_block = (b[4] or "").strip()
                if text_block:
                    parts.append(text_block)

        # NOTE: Uncomment ONLY ONE of the above sections.
        # If neither is uncommented, parts will be empty.
        pass

    finally:
        doc.close()

    duration = time.time() - start
    print(f"Extraction time: {duration:.2f} seconds")

    return ("\n".join(parts))

def html_to_docx_bytes(html: str) -> bytes:
    """
    Convert HTML (like your IM template) to a reasonably structured DOCX.
    - h1/h2/h3 -> Word headings
    - p       -> paragraphs
    - ul/ol   -> bullet/numbered lists
    - table   -> grid table with header row
    """
    soup = BeautifulSoup(html, "html.parser")
    document = Document()

    # Optional: add main title from <title> if present
    if soup.title and soup.title.string:
        title_para = document.add_paragraph(soup.title.string.strip())
        title_para.style = "Title"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Work on <body> if present, else whole soup
    root: Tag = soup.body if soup.body else soup

    def handle_element(el: Tag, doc: Document):
        # Skip non-tag nodes handled via parent
        if isinstance(el, NavigableString):
            text = el.strip()
            if text:
                doc.add_paragraph(text)
            return

        # Headings
        if el.name in ["h1", "h2", "h3"]:
            text = el.get_text(strip=True)
            if not text:
                return
            level = {"h1": 1, "h2": 2, "h3": 3}[el.name]
            doc.add_heading(text, level=level)
            return

        # Paragraphs
        if el.name == "p":
            text = el.get_text(" ", strip=True)
            if text:
                doc.add_paragraph(text)
            return

        # Unordered list (bullets)
        if el.name == "ul":
            for li in el.find_all("li", recursive=False):
                text = li.get_text(" ", strip=True)
                if text:
                    p = doc.add_paragraph(text)
                    p.style = "List Bullet"
            return

        # Ordered list (numbers)
        if el.name == "ol":
            for li in el.find_all("li", recursive=False):
                text = li.get_text(" ", strip=True)
                if text:
                    p = doc.add_paragraph(text)
                    p.style = "List Number"
            return

        # Tables
        if el.name == "table":
            rows = el.find_all("tr")
            if not rows:
                return

            # Determine max columns from first row with cells
            max_cols = 0
            row_cells = []
            for r in rows:
                cells = r.find_all(["th", "td"], recursive=False)
                if cells:
                    row_cells.append(cells)
                    max_cols = max(max_cols, len(cells))

            if max_cols == 0:
                return

            table = doc.add_table(rows=len(row_cells), cols=max_cols)
            table.style = "Table Grid"

            for r_idx, cells in enumerate(row_cells):
                row = table.rows[r_idx]
                for c_idx in range(max_cols):
                    cell = row.cells[c_idx]
                    if c_idx < len(cells):
                        text = cells[c_idx].get_text(" ", strip=True)
                        cell.text = text

            return

        # Structural wrappers: handle children
        if el.name in ["div", "section", "main", "header", "footer"]:
            for child in el.children:
                if isinstance(child, (Tag, NavigableString)):
                    handle_element(child, doc)
            return

        # Fallback: just process children
        for child in el.children:
            if isinstance(child, (Tag, NavigableString)):
                handle_element(child, doc)

    # Process top-level children of root
    for child in root.children:
        if isinstance(child, (Tag, NavigableString)):
            handle_element(child, document)

    # Save to bytes
    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()

@app.post('/html-to-docx')
async def html_to_docx(
    request: Request,
    html: Optional[str] = Form(None),
    file: Optional[UploadFile] = Form(None)
):
    try:
        file_name = "output.docx"
        if request.headers.get("content-type","").startswith("application/json"):
            payload = HtmlToDocxPayload(**await request.json())
            html_content = payload.html
            file_name = payload.file_name
        elif file:
            html_content = (await file.read()).decode("utf-8")
        elif html:
            html_content = html
        else:
            html_content = (await request.body()).decode("utf-8")

        if not html_content.strip():
            return JSONResponse(
                status_code = 400,
                content = {"error":"HTML content is empty"},
            )
        docx_bytes = html_to_docx_bytes(html_content)

        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type= "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition":f'attachment;filename="{file_name}"'
            },
        )


            
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error":str(e)}
        )
